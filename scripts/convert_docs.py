#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Convert Fintech Payment OpenAPI docx files to structured HTML documentation.
UI/UX inspired by Stripe API documentation (https://docs.stripe.com/api).
"""

import os
import re
import sys
import json
import html
import shutil
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from opencc import OpenCC

# OpenCC instance for simplified -> traditional conversion
_cc = OpenCC('s2t')

# Configuration
INPUT_DIR = "/Users/susi/Documents/Project/ftp.openapi.trae/Fintech Payment Platform OpenAPI說明文檔20260304"
OUTPUT_DIR = "/Users/susi/Documents/Project/ftp.openapi.trae/docs"

# XML namespaces for image extraction
_NSW = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
_NSA = 'http://schemas.openxmlformats.org/drawingml/2006/main'
_NSR = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'

# ===========================================================================
# Category definitions (new structure per user requirements)
# ===========================================================================
CATEGORIES = {
    'card-online':  {'label': 'Card Pay（Online）',  'icon': '💳', 'desc': '支持 Visa、Mastercard、銀聯等卡組織的線上卡支付', 'color': 'card-pay'},
    'card-offline': {'label': 'Card Pay（Offline）', 'icon': '💳', 'desc': '支持 Visa、Mastercard 等卡組織的線下物理卡支付',     'color': 'card-pay'},
    'qr-online':    {'label': 'QR Pay（Online）',    'icon': '📱', 'desc': '支持 Alipay、WeChat 等行動支付的線上掃碼支付',       'color': 'qr-pay'},
    'qr-offline':   {'label': 'QR Pay（Offline）',   'icon': '📱', 'desc': '支持付款碼、CashAPI 等線下掃碼支付',                  'color': 'qr-pay'},
    'other':        {'label': 'Other Pay',           'icon': '⚙️', 'desc': '外幣支付等特殊場景支付方案',                          'color': 'other-pay'},
}

# Category display order
CATEGORY_ORDER = ['card-online', 'card-offline', 'qr-online', 'qr-offline', 'other']

# Document mapping (ordered by category, then brand/product)
DOCS = [
    {
        "file": "3.1_卡支付（VMC）：Global Payments Online無卡支付接口文檔_V1.0（托管）.docx",
        "id": "global-payments-online",
        "category": "card-online",
        "title": "Global Payments Online 無卡支付",
        "subtitle": "Card Pay（Online） - 託管",
        "status": "active"
    },
    {
        "file": "4.5_卡支付（VMC）：Sinopay SPGS Online支付接口文檔_V1.0.docx",
        "id": "fintech-payment-online",
        "category": "card-online",
        "title": "Fintech Payment Online 支付",
        "subtitle": "Card Pay（Online）",
        "status": "upcoming"
    },
    {
        "file": "3.2_卡支付（VMC）：Global Payments Offline物理卡支付接口文檔_V1.0（備註：未提供任何人對接）.docx",
        "id": "global-payments-offline",
        "category": "card-offline",
        "title": "Global Payments Offline 物理卡支付",
        "subtitle": "Card Pay（Offline）",
        "status": "upcoming"
    },
    {
        "file": "3.3_線上支付：Alipay WAP支付接口文檔_V1.0.docx",
        "id": "alipay-wap",
        "category": "qr-online",
        "title": "Alipay WAP 支付",
        "subtitle": "QR Pay（Online）",
        "status": "active"
    },
    {
        "file": "3.4_線上支付：Alipay WEB支付接口文檔_V1.0.docx",
        "id": "alipay-web",
        "category": "qr-online",
        "title": "Alipay WEB 支付",
        "subtitle": "QR Pay（Online）",
        "status": "active"
    },
    {
        "file": "4.2_線上支付：Alipay APP支付接口文檔_V1.0.docx",
        "id": "alipay-app",
        "category": "qr-online",
        "title": "Alipay APP 支付",
        "subtitle": "QR Pay（Online）",
        "status": "active"
    },
    {
        "file": "4.3_線上支付：Alipay Plus+支付接口文檔_V1.0.docx",
        "id": "alipay-plus",
        "category": "qr-online",
        "title": "Alipay Plus+ 支付",
        "subtitle": "QR Pay（Online）",
        "status": "active"
    },
    {
        "file": "3.5_線上支付：Wechat WEB支付接口文檔_V1.0.docx",
        "id": "wechat-web",
        "category": "qr-online",
        "title": "Wechat WEB 支付",
        "subtitle": "QR Pay（Online）",
        "status": "active"
    },
    {
        "file": "3.8_線上支付：Wechat WAP支付接口文檔_V1.0.docx",
        "id": "wechat-wap",
        "category": "qr-online",
        "title": "Wechat WAP 支付",
        "subtitle": "QR Pay（Online）",
        "status": "active"
    },
    {
        "file": "3.9_線上支付：Wechat 公眾帳號和小程序支付接口文檔_V1.0.docx",
        "id": "wechat-miniprogram",
        "category": "qr-online",
        "title": "Wechat 公眾帳號和小程序支付",
        "subtitle": "QR Pay（Online）",
        "status": "active"
    },
    {
        "file": "4.4_線上支付：Wechat APP支付接口文檔_V1.0.docx",
        "id": "wechat-app",
        "category": "qr-online",
        "title": "Wechat APP 支付",
        "subtitle": "QR Pay（Online）",
        "status": "active"
    },
    {
        "file": "3.6_線上支付：掃碼支付接口文檔_V1.0.docx",
        "id": "qr-code",
        "category": "qr-online",
        "title": "掃碼支付",
        "subtitle": "QR Pay（Online）",
        "status": "active"
    },
    {
        "file": "3.7_線下支付：付款碼支付接口文檔_V1.0.docx",
        "id": "barcode",
        "category": "qr-offline",
        "title": "付款碼支付",
        "subtitle": "QR Pay（Offline）",
        "status": "active"
    },
    {
        "file": "4.1_線下支付：CashAPI接口文檔（招財貓系統專用）_V1.0.docx",
        "id": "cash-api",
        "category": "qr-offline",
        "title": "CashAPI 支付",
        "subtitle": "QR Pay（Offline） - 招財貓系統專用",
        "status": "active"
    },
    {
        "file": "4.0_特殊支付：FTP外幣支付接口文檔（備註：PROD未啟用，僅當初申請HK文件所用）.docx",
        "id": "ftp-foreign",
        "category": "other",
        "title": "FTP 外幣支付",
        "subtitle": "Other Pay",
        "status": "upcoming"
    },
]


# ===========================================================================
# Text processing helpers
# ===========================================================================

def to_traditional(text):
    """Convert simplified Chinese to traditional Chinese."""
    if not text:
        return text
    return _cc.convert(text)


def replace_terms(text):
    """Replace brand terms according to requirements.

    - Sinopay / 大金  -> Fintech Payment
    - SP / 威富通     -> Swiftpass
    """
    if not text:
        return text
    # Sinopay -> Fintech Payment (case-insensitive, whole word)
    text = re.sub(r'(?i)\bsinopay\b', 'Fintech Payment', text)
    # 大金 -> Fintech Payment
    text = text.replace('大金', 'Fintech Payment')
    # 威富通 -> Swiftpass
    text = text.replace('威富通', 'Swiftpass')
    # SP -> Swiftpass (only as standalone word, not part of SPGS etc.)
    text = re.sub(r'\bSP\b', 'Swiftpass', text)
    return text


def clean_text(text):
    """Apply all text cleaning: brand replacement + traditional conversion."""
    text = replace_terms(text)
    text = to_traditional(text)
    return text


def extract_images_from_para(para, doc, doc_id, img_counter):
    """Extract images embedded in a paragraph's XML."""
    images = []
    drawings = para._element.findall(f'.//{{{_NSW}}}drawing')
    for drawing in drawings:
        blips = drawing.findall(f'.//{{{_NSA}}}blip')
        for blip in blips:
            embed = blip.get(f'{{{_NSR}}}embed')
            if embed and embed in doc.part.rels:
                rel = doc.part.rels[embed]
                image_data = rel.target_part.blob
                content_type = rel.target_part.content_type
                ext = 'png'
                if 'jpeg' in content_type or 'jpg' in content_type:
                    ext = 'jpg'
                elif 'png' in content_type:
                    ext = 'png'
                elif 'gif' in content_type:
                    ext = 'gif'
                img_counter[0] += 1
                img_name = f"{doc_id}_{img_counter[0]}.{ext}"
                img_dir = os.path.join(OUTPUT_DIR, 'assets', 'images')
                img_path = os.path.join(img_dir, img_name)
                os.makedirs(img_dir, exist_ok=True)
                with open(img_path, 'wb') as f:
                    f.write(image_data)
                images.append(img_name)
    return images


def slugify(text):
    """Convert text to URL-friendly slug."""
    text = re.sub(r'[\s\uff08\uff09\uff1a\uff0f\u3001]+', '-', text)
    text = re.sub(r'[^\w\-]', '', text)
    text = re.sub(r'-+', '-', text)
    return text.strip('-').lower()


def is_json_block(text):
    """Check if text looks like JSON."""
    stripped = text.strip()
    return (stripped.startswith('{') and stripped.endswith('}')) or \
           (stripped.startswith('[') and stripped.endswith(']')) or \
           ('"' in stripped and ':' in stripped and ('{' in stripped or '[' in stripped))


def is_xml_block(text):
    """Check if text looks like XML/HTML."""
    stripped = text.strip()
    return stripped.startswith('<') and stripped.endswith('>') and ('</' in stripped or '/>' in stripped)


def detect_code_type(text):
    """Detect if text is code/json/xml."""
    text = text.strip()
    if is_json_block(text):
        return 'json'
    if is_xml_block(text):
        return 'xml'
    if text.startswith('HTTP/') or text.startswith('GET ') or text.startswith('POST '):
        return 'http'
    return None


def format_code_block(text, lang=None):
    """Format text as a code block."""
    if lang:
        return f'<pre><span class="code-label">{lang.upper()}</span><code class="language-{lang}">{html.escape(text)}</code></pre>'
    return f'<pre><code>{html.escape(text)}</code></pre>'


def process_paragraph_text(para):
    """Extract text with basic formatting from a paragraph."""
    texts = []
    for run in para.runs:
        text = run.text
        if not text:
            continue
        text = clean_text(text)
        if run.bold:
            text = f'<strong>{html.escape(text)}</strong>'
        else:
            text = html.escape(text)
        texts.append(text)
    return ''.join(texts)


def extract_docx_content(doc_path, doc_id=''):
    """Extract structured content from a docx file."""
    doc = Document(doc_path)
    content = []
    img_counter = [0]

    for element in doc.element.body:
        tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag

        if tag == 'p':
            for para in doc.paragraphs:
                if para._element is element:
                    text = para.text.strip()

                    # Extract images first (images may accompany text or be standalone)
                    images = extract_images_from_para(para, doc, doc_id, img_counter)
                    for img_name in images:
                        content.append({'type': 'image', 'src': img_name})

                    if not text:
                        continue

                    # Clean text (brand replacement + traditional Chinese)
                    text = clean_text(text)

                    style = para.style.name if para.style else 'Normal'

                    if 'toc' in style.lower():
                        continue

                    # Skip TOC entries (lines that look like "title\tpage_number")
                    if re.match(r'^.+\t\d+$', text):
                        continue

                    if style.startswith('Heading'):
                        level = int(style.replace('Heading ', '')) if style.replace('Heading ', '').isdigit() else 1
                        content.append({
                            'type': 'heading',
                            'level': level,
                            'text': text
                        })
                    elif is_json_block(text) or is_xml_block(text):
                        lang = detect_code_type(text)
                        content.append({
                            'type': 'code',
                            'lang': lang or 'text',
                            'text': text
                        })
                    else:
                        content.append({
                            'type': 'paragraph',
                            'text': text
                        })
                    break

        elif tag == 'tbl':
            for tbl in doc.tables:
                if tbl._element is element:
                    table_data = []
                    for row in tbl.rows:
                        row_data = []
                        for cell in row.cells:
                            cell_text = clean_text(cell.text.strip())
                            row_data.append(cell_text)
                        table_data.append(row_data)

                    if table_data:
                        content.append({
                            'type': 'table',
                            'data': table_data
                        })
                    break

    return content


def render_html(content, doc_info, is_common=False, depth=2):
    """Render structured content to HTML."""
    html_parts = []
    toc_items = []

    # Page title
    html_parts.append(f'<h1 id="overview">{html.escape(doc_info["title"])}')
    if doc_info.get('subtitle'):
        html_parts.append(f'<span class="subtitle">{html.escape(doc_info["subtitle"])}</span>')
    if doc_info.get('status') == 'upcoming':
        html_parts.append(f'<span class="status-badge upcoming">暫未開放</span>')
    html_parts.append('</h1>')

    # Meta info
    html_parts.append('<div class="page-meta">')
    html_parts.append(f'<span class="page-meta-item">文檔版本：V1.0</span>')
    html_parts.append(f'<span class="page-meta-item">接口版本：V1.0</span>')
    html_parts.append('</div>')

    # Upcoming warning
    if doc_info.get('status') == 'upcoming':
        html_parts.append('<div class="callout warning">')
        html_parts.append('<div class="callout-title">⚠️ 提示</div>')
        html_parts.append('本接口暫未開放使用，僅供參考。具體上線時間請聯繫平臺客服。')
        html_parts.append('</div>')

    for item in content:
        if item['type'] == 'heading':
            level = item['level']
            text = item['text']
            slug = slugify(text)

            if text == '目錄' or 'docx' in text.lower() or text in [doc_info['title'], doc_info.get('subtitle', '')]:
                continue

            if level == 2:
                toc_items.append({'level': 2, 'text': text, 'slug': slug})
            elif level == 3:
                toc_items.append({'level': 3, 'text': text, 'slug': slug})

            html_parts.append(f'<h{level} id="{slug}">{html.escape(text)}</h{level}>')

        elif item['type'] == 'paragraph':
            text = item['text']
            if not text or text in ['文檔版本：1.0', '接口版本：1.0', '修改日誌']:
                continue
            html_parts.append(f'<p>{html.escape(text)}</p>')

        elif item['type'] == 'code':
            lang = item.get('lang', 'text')
            text = item['text']
            html_parts.append(format_code_block(text, lang))

        elif item['type'] == 'image':
            img_prefix = '../' * depth
            html_parts.append(f'<div class="doc-image"><img src="{img_prefix}assets/images/{item["src"]}" alt="文檔圖片" loading="lazy"></div>')

        elif item['type'] == 'table':
            data = item['data']
            if not data:
                continue

            html_parts.append('<div class="table-wrapper">')
            html_parts.append('<table>')

            html_parts.append('<thead><tr>')
            for cell in data[0]:
                html_parts.append(f'<th>{html.escape(cell)}</th>')
            html_parts.append('</tr></thead>')

            html_parts.append('<tbody>')
            for row in data[1:]:
                html_parts.append('<tr>')
                for cell in row:
                    cell_html = html.escape(cell)
                    if cell == 'M' or cell == 'O' or cell == 'C':
                        if cell == 'M':
                            cell_html = '<span class="param-required yes">必填</span>'
                        elif cell == 'O':
                            cell_html = '<span class="param-required no">選填</span>'
                        elif cell == 'C':
                            cell_html = '<span class="param-required no">條件</span>'
                    html_parts.append(f'<td>{cell_html}</td>')
                html_parts.append('</tr>')
            html_parts.append('</tbody>')
            html_parts.append('</table>')
            html_parts.append('</div>')

    return '\n'.join(html_parts), toc_items


# ===========================================================================
# HTML template (Stripe-inspired)
# ===========================================================================

def generate_full_html(content_html, toc_items, doc_info, depth=0):
    """Generate complete HTML page with template."""
    prefix = '../' * depth

    # Build TOC HTML
    toc_html = ''
    if toc_items:
        toc_html = '<nav class="toc-nav">\n'
        toc_html += '<div class="toc-nav-title">本頁目錄</div>\n'
        toc_html += '<ul>\n'
        for item in toc_items:
            cls = f'toc-h{item["level"]}'
            toc_html += f'  <li><a href="#{item["slug"]}" class="{cls}">{html.escape(item["text"])}</a></li>\n'
        toc_html += '</ul>\n</nav>'

    category = doc_info.get('category', '')
    page_id = doc_info.get('id', '')

    nav_html = generate_sidebar(category, page_id, depth)

    # First product doc link for nav
    first_doc = DOCS[0]

    return f'''<!DOCTYPE html>
<html lang="zh-Hant">
<head>
  <meta charset="UTF-8">
  <meta name="viewport" content="width=device-width, initial-scale=1.0">
  <title>{html.escape(doc_info['title'])} | Fintech Payment OpenAPI 文檔</title>
  <link rel="stylesheet" href="{prefix}assets/css/style.css">
</head>
<body>
  <header class="site-header">
    <div class="header-inner">
      <a href="{prefix}index.html" class="logo">
        <span class="logo-icon">FP</span>
        <span class="logo-text">Fintech Payment <span class="logo-sub">OpenAPI</span></span>
      </a>
      <nav class="header-nav">
        <a href="{prefix}index.html">首頁</a>
        <a href="{prefix}common/quick-start.html">接入指引</a>
        <a href="{prefix}common/data-format.html">通用規則</a>
        <a href="{prefix}payments/{first_doc["category"]}/{first_doc["id"]}.html">產品文檔</a>
      </nav>
      <button class="menu-toggle" aria-label="選單">☰</button>
    </div>
  </header>

  <div class="docs-layout">
    {nav_html}

    <main class="main-content">
      {content_html}
    </main>

    {toc_html}
  </div>

  <script src="{prefix}assets/js/main.js"></script>
</body>
</html>'''


def generate_sidebar(active_category, active_page, depth=0):
    """Generate sidebar navigation HTML."""
    prefix = '../' * depth

    # Group docs by category
    grouped = {}
    for d in DOCS:
        cat = d['category']
        if cat not in grouped:
            grouped[cat] = []
        grouped[cat].append(d)

    parts = ['<aside class="sidebar">\n']

    # Common rules section
    parts.append('<div class="sidebar-group" data-group="common">\n')
    parts.append('  <div class="sidebar-group-title"><span class="toggle-icon">▼</span>通用規則</div>\n')
    parts.append('  <ul class="sidebar-links">\n')
    for link_id, link_label, link_path in [
        ('quick-start', '快速入門', f'{prefix}common/quick-start.html'),
        ('endpoints', 'API 調用地址', f'{prefix}common/endpoints.html'),
        ('data-format', '數據格式', f'{prefix}common/data-format.html'),
        ('signature', '數字簽名', f'{prefix}common/signature.html'),
        ('request-headers', '請求頭說明', f'{prefix}common/request-headers.html'),
        ('error-codes', '錯誤碼說明', f'{prefix}common/error-codes.html'),
        ('notification', '支付通知說明', f'{prefix}common/notification.html'),
    ]:
        active = ' active' if active_page == link_id else ''
        parts.append(f'    <li><a href="{link_path}" class="sub-link{active}">{link_label}</a></li>\n')
    parts.append('  </ul>\n</div>\n')

    # Payment categories (new structure)
    for cat_key in CATEGORY_ORDER:
        if cat_key not in grouped:
            continue
        cat_info = CATEGORIES[cat_key]
        docs_list = grouped[cat_key]

        parts.append(f'<div class="sidebar-group" data-group="{cat_key}">\n')
        parts.append(f'  <div class="sidebar-group-title"><span class="toggle-icon">▼</span>{cat_info["icon"]} {cat_info["label"]}</div>\n')
        parts.append('  <ul class="sidebar-links">\n')

        for d in docs_list:
            is_active = active_page == d['id']
            active_cls = ' active' if is_active else ''
            status_tag = ''
            if d['status'] == 'upcoming':
                status_tag = ' <span class="status-tag upcoming">暫未開放</span>'
            parts.append(f'    <li><a href="{prefix}payments/{cat_key}/{d["id"]}.html" class="sub-link{active_cls}">{html.escape(d["title"])}{status_tag}</a></li>\n')

        parts.append('  </ul>\n</div>\n')

    parts.append('</aside>\n')
    return ''.join(parts)


def build_document(doc_info):
    """Build a single document page."""
    doc_path = os.path.join(INPUT_DIR, doc_info['file'])
    if not os.path.exists(doc_path):
        print(f"  ⚠️ 文件不存在: {doc_path}")
        return False

    print(f"  處理: {doc_info['title']}")
    content = extract_docx_content(doc_path, doc_info['id'])
    content_html, toc_items = render_html(content, doc_info, depth=2)

    cat = doc_info['category']
    output_path = os.path.join(OUTPUT_DIR, 'payments', cat, f"{doc_info['id']}.html")
    depth = 2  # payments/cat/file.html

    full_html = generate_full_html(content_html, toc_items, doc_info, depth)

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(full_html)

    return True


def _write_common_page(page_id, title, content_html, toc_items=None):
    """Write a common page to the output directory."""
    if toc_items is None:
        toc_items = []
    doc_info = {'title': title, 'subtitle': '通用規則', 'status': 'active', 'id': page_id, 'category': ''}
    full_html = generate_full_html(content_html, toc_items, doc_info, depth=1)
    output_path = os.path.join(OUTPUT_DIR, 'common', f"{page_id}.html")
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(full_html)
    print(f"  建構: {title}")


# Interface support matrix: (doc_id, display_name, [下單, 調起支付, 查詢, 退款, 退款查詢, 支付通知])
_INTERFACE_MATRIX = [
    ('global-payments-online',  'GP Online 無卡支付',      ['✅', '✅', '✅', '—', '—', '✅']),
    ('global-payments-offline', 'GP Offline 物理卡支付',   ['✅', '—', '✅', '—', '—', '—']),
    ('fintech-payment-online',  'Fintech Payment Online',  ['✅', '—', '✅', '✅', '✅', '✅']),
    ('alipay-wap',              'Alipay WAP',              ['✅', '✅', '✅', '✅', '✅', '✅']),
    ('alipay-web',              'Alipay WEB',              ['✅', '✅', '✅', '✅', '✅', '✅']),
    ('alipay-app',              'Alipay APP',              ['✅', '—', '✅', '✅', '✅', '✅']),
    ('alipay-plus',             'Alipay Plus+',            ['✅', '—', '✅', '✅', '✅', '✅']),
    ('wechat-web',              'Wechat WEB',              ['✅', '✅', '✅', '✅', '✅', '✅']),
    ('wechat-wap',              'Wechat WAP',              ['✅', '—', '✅', '✅', '✅', '✅']),
    ('wechat-app',              'Wechat APP',              ['✅', '—', '✅', '✅', '✅', '✅']),
    ('wechat-miniprogram',      'Wechat 公眾帳號/小程序',  ['✅', '✅', '✅', '✅', '✅', '✅']),
    ('qr-code',                 '掃碼支付',                ['✅', '✅', '✅', '✅', '✅', '✅']),
    ('barcode',                 '付款碼支付',              ['✅', '—', '✅', '✅', '✅', '—']),
    ('cash-api',                'CashAPI',                 ['✅', '—', '—', '✅', '—', '—']),
    ('ftp-foreign',             'FTP 外幣支付',            ['✅', '—', '✅', '—', '—', '—']),
]

# Common error codes (appear in all 15 docs)
_COMMON_ERRORS = [
    ('MISSING_FIELD', '字段未傳入'), ('INVALID_FIELD', '字段校驗錯誤'),
    ('NO_EXISTS', '數據不存在'), ('OTHER', '其他錯誤'),
    ('INVALID_SCOPE', '數據超出範圍'), ('INVALID_LENGTH', '數據長度超過限制'),
    ('INCORRECT_SIGNATURE', '無效簽名'), ('EXPIRED_ORDER', '訂單已過期'),
    ('PAY_GATEWAY_FAIL', '支付網關異常'), ('INVALID_SERVICE', '非法接口類型'),
    ('INVALID_CLIENT_ID', '無效客戶端ID'), ('INVALID_STORE_ID', '無效門店ID'),
    ('INVALID_AMOUNT', '交易金額不正確'), ('INVALID_ORDER_STATUS', '無效的訂單狀態'),
    ('INVALID_STATUS', '非法狀態'), ('INVALID_REFUND_FEE', '退款金額不正確'),
    ('LACK_PARAMS', '缺少參數'), ('REQUIRE_POST_METHOD', '請使用POST方法'),
    ('POST_DATA_EMPTY', 'POST數據為空'), ('NOT_UTF8', '編碼格式錯誤'),
    ('URL_FORMAT_ERROR', 'URL格式不正確'), ('IP_FORMAT_ERROR', 'IP格式不正確'),
    ('XML_FORMAT_ERROR', 'XML格式錯誤'), ('DATE_FORMAT_ERROR', '日期格式不正確'),
    ('NO_AUTH', '權限不足'), ('ACCESS_DENIED', '拒絕訪問'),
    ('ACCOUNT_LOCKED', '商戶或門店的賬號被鎖定'), ('SN_IS_EXISTS', '編號已存在'),
    ('TRADE_NOT_EXIST', '交易不存在'), ('TRADE_FAILED', '交易失敗'),
    ('TRADE_TIMEOUT', '交易超時'), ('TRADE_HAS_SUCCESS', '交易已被支付'),
    ('REFUND_FAILED', '退款失敗'), ('DEDUCTIONS_FAILED', '扣款失敗'),
    ('REVERSE_FAILED', '撤銷授權失敗'), ('NOTIFY_FAIL', '通知異常'),
    ('ORDER_CLOSED', '訂單已關閉'), ('DATA_NOT_EXIST', '查詢的數據不存在'),
    ('UNSUPPORTED_SERVICE', '不支持的接口類型'), ('VALID_FLOW_FAIL', '請求過於頻繁'),
    ('UNKNOWN_ERROR', '未知錯誤'),
]

# Card-specific error codes (only in card payment docs)
_CARD_ERRORS = [
    ('DECLINED', '拒絕處理'), ('GENERAL_DECLINE', '拒絕處理'),
    ('INVALID_CARD', '無效帳號'), ('INVALID_CARD_TYPE', '卡類型無效'),
    ('CARD_TYPE_NOT_ACCEPTED', '不支持卡類型'), ('EXPIRED_CARD', '已過期的卡'),
    ('INVALID_CVN', '無效CVN'), ('CVN_NOT_MATCH', 'CVN不匹配'),
    ('CV_FAILED', 'CVN檢查被拒絕'), ('AVS_FAILED', 'AVS檢查被拒絕'),
    ('INSUFFICIENT_FUND', '帳戶資金不足'), ('EXCEEDS_CREDIT_LIMIT', '該卡已達到信用額度'),
    ('STOLEN_LOST_CARD', '卡失竊或遺失'), ('UNAUTHORIZED_CARD', '無效的卡或未經授權的無卡交易'),
    ('PROCESSOR_DECLINED', '發卡行拒絕支付'), ('PROCESSOR_TIMEOUT', '支付處理超時'),
    ('PROCESSOR_UNAVAILABLE', '銀行故障'), ('ISSUER_UNAVAILABLE', '發卡行不可用'),
    ('CONTACT_PROCESSOR', '發卡銀行對請求有疑問'), ('AUTHENTICATION_FAILED', '身份驗證失敗'),
    ('CONSUMER_AUTHENTICATION_FAILED', '付款人無法驗證'), ('CONSUMER_AUTHENTICATION_REQUIRED', '未驗證身份信息'),
    ('PENDING_AUTHENTICATION', '等待身份驗證'), ('PENDING_REVIEW', '等待審查'),
    ('AUTHORIZED_PENDING_REVIEW', '待授權處理審查'), ('AUTHORIZED_RISK_DECLINED', '拒絕風險授權'),
    ('AUTH_ALREADY_REVERSED', '授權已經撤銷'), ('CAPTURE_ALREADY_VOIDED', '扣款已失效'),
    ('NOT_VOIDABLE', '扣款無效'), ('VOIDED', '已廢棄'),
    ('TRANSACTION_ALREADY_SETTLED', '交易已結算'), ('TRANSACTION_ALREADY_REVERSED_OR_SETTLED', '交易已經結算或撤銷'),
    ('EXCEEDS_AUTH_AMOUNT', '金額不正確'), ('INVALID_ACCOUNT', '無效的帳號'),
    ('SUSPENDED_ACCOUNT', '帳戶被凍結'), ('BLACKLISTED_CUSTOMER', '訂單信息不匹配'),
    ('CUSTOMER_WATCHLIST_MATCH', '客戶實體信息被限制'), ('IP_COUNTRY_WATCHLIST_MATCH', 'IP被限制'),
    ('EMAIL_COUNTRY_WATCHLIST_MATCH', '電子郵件被限制'), ('ADDRESS_COUNTRY_WATCHLIST_MATCH', '訂單被限制'),
    ('SCORE_EXCEEDS_THRESHOLD', '安全檢查不通過'), ('DECISION_PROFILE_REJECT', '卡審核失敗'),
    ('DECISION_PROFILE_REVIEW', '卡待審核'), ('DECISION_PROFILE_CHALLENGE', '存在風險交易'),
    ('CHALLENGE', '風險交易'), ('PAYMENT_REFUSED', '拒絕支付'),
    ('REJECTED', '已拒絕'), ('DUPLICATE_REQUEST', '重複請求'),
    ('INVALID_REQUEST', '非法請求'), ('INVALID_DATA', '無效數據'),
    ('INVALID_PAYMENT_ID', '請求ID無效'), ('INVALID_MERCHANT_CONFIGURATION', '商戶配置異常'),
    ('INVALID_OR_MISSING_CONFIG', '商戶配置異常'), ('MISSING_AUTH', '未授權異常'),
    ('SERVER_ERROR', '服務錯誤'), ('SERVER_TIMEOUT', '服務響應超時'),
    ('SERVICE_TIMEOUT', '服務響應超時'), ('SYSTEM_ERROR', '系統故障'),
    ('ROUTE_SERVICE_NOT_FOUND', '路由服務標識無效'), ('UNSUPPORTED_CHARACTER_SET', '地址驗證服務不支持'),
    ('BOLETO_DECLINED', 'boleto請求被拒絕'), ('DECLINED_CHECK', '支付被拒絕'),
    ('DEBIT_CARD_USAGE_LIMIT_EXCEEDED', '卡使用次數過多'), ('ACCOUNT_NOT_ALLOWED_CREDIT', '退款被禁止'),
    ('INSUFFICIENT_ADDRESS_INFORMATION', '地址信息不足'), ('UNVERIFIABLE_ADDRESS', '地址無法驗證'),
    ('ADDRESS_MATCH_NOT_FOUND', '地址無法匹配'), ('MULTIPLE_ADDRESS_MATCHES', '存在多個地址匹配'),
    ('MULTIPLE_ADDRESS_MATCHES_INTERNATIONAL', '存在多個地址情況'), ('APARTMENT_NUMBER_NOT_FOUND', '地址信息不正確'),
    ('HOUSE_OR_BOX_NUMBER_NOT_FOUND', '門牌號或郵局不正確'), ('BOX_NUMBER_NOT_FOUND', '郵政信息無效'),
    ('STREET_NAME_NOT_FOUND', '郵政編碼與街道不對應'), ('POSTAL_CODE_NOT_FOUND', '郵政編碼無效'),
]


def build_common_pages():
    """Build all 7 common rule pages (hand-crafted for clarity)."""
    print("\n🔧 建構通用規則頁面...")

    # 1. Quick Start
    toc = [
        {'level': 2, 'text': '接入前準備', 'slug': 'prerequisites'},
        {'level': 2, 'text': '支付流程', 'slug': 'payment-flow'},
        {'level': 2, 'text': '快速接入步驟', 'slug': 'quick-steps'},
        {'level': 2, 'text': '完整代碼示例', 'slug': 'code-example'},
    ]
    content = '''<h1 id="overview">快速入門</h1>
<div class="page-meta">
  <span class="page-meta-item">文檔版本：V1.0</span>
</div>
<p>本指南將引導您從零開始完成 Fintech Payment OpenAPI 的接入，包括環境準備、簽名計算、發起支付請求及處理支付通知。</p>

<h2 id="prerequisites">接入前準備</h2>
<p>在開始接入前，請確保您已具備以下條件：</p>
<div class="table-wrapper">
<table>
<thead><tr><th>項目</th><th>說明</th><th>獲取方式</th></tr></thead>
<tbody>
<tr><td>商戶帳號</td><td>Fintech Payment 平台商戶帳號</td><td>由平台商務團隊開通</td></tr>
<tr><td>client_id</td><td>客戶端 ID，標識商戶身份</td><td>商戶後台或商務團隊提供</td></tr>
<tr><td>store_id</td><td>門店 ID，標識具體門店</td><td>商戶後台或商務團隊提供</td></tr>
<tr><td>商戶秘鑰</td><td>用於數字簽名的通信金鑰</td><td>商務團隊提供，請妥善保管</td></tr>
<tr><td>服務器環境</td><td>支持 HTTPS POST 請求的服務器</td><td>商戶自行準備</td></tr>
</tbody>
</table>
</div>

<h2 id="payment-flow">支付流程</h2>
<p>典型的支付流程如下：</p>
<div class="callout info">
  <div class="callout-title">📋 業務流程</div>
  <p><strong>1. 下單</strong> → 商戶服務器調用 <code>/pay</code> 接口創建訂單<br>
  <strong>2. 調起支付</strong> → 商戶根據返回的支付鏈接/二維碼引導用戶支付<br>
  <strong>3. 支付通知</strong> → 平台異步通知商戶 <code>notify_url</code> 支付結果<br>
  <strong>4. 訂單查詢</strong> → 商戶主動調用 <code>/query</code> 確認訂單狀態<br>
  <strong>5. 退款</strong>（可選）→ 商戶調用 <code>/refund</code> 發起退款</p>
</div>
<p>所有接口調用都必須在商戶服務器端完成，切勿在客戶端直接調用。</p>

<h2 id="quick-steps">快速接入步驟</h2>

<h3>步驟一：構建請求參數</h3>
<p>組裝 JSON 格式的請求參數，包含業務參數和公共參數：</p>
<pre><span class="code-label">JSON</span><code class="language-json">{
  "service": "pay.weixin.web.intl",
  "client_id": "1001",
  "store_id": "1001001",
  "order_no": "ORDER20260708001",
  "amount": 1.00,
  "body": "測試商品",
  "notify_url": "https://your-domain.com/notify",
  "callback_url": "https://your-domain.com/result",
  "client_ip": "127.0.0.1",
  "nonce_str": "abc123def456",
  "lang": "en_US"
}</code></pre>

<h3>步驟二：計算簽名</h3>
<p>將所有參數（除 <code>sign</code> 和 <code>lang</code> 外）按 ASCII 碼從小到大排序，拼接成 QueryString 格式，再拼接商戶秘鑰後做 MD5 運算：</p>
<pre><span class="code-label">PYTHON</span><code class="language-python">import hashlib

params = {
    "service": "pay.weixin.web.intl",
    "client_id": "1001",
    "store_id": "1001001",
    "order_no": "ORDER20260708001",
    "amount": "1.00",
    "body": "測試商品",
    "notify_url": "https://your-domain.com/notify",
    "callback_url": "https://your-domain.com/result",
    "client_ip": "127.0.0.1",
    "nonce_str": "abc123def456",
}

# 1. 按 key 的 ASCII 碼從小到大排序
sorted_params = sorted(params.items())

# 2. 拼接成 key=value&key=value 格式（空值不傳遞）
sign_str = "&".join(f"{k}={v}" for k, v in sorted_params if v)

# 3. 拼接商戶秘鑰
merchant_key = "YOUR_MERCHANT_KEY"
sign_str += f"&key={merchant_key}"

# 4. MD5 運算並轉大寫
sign = hashlib.md5(sign_str.encode("utf-8")).hexdigest().upper()
print(sign)</code></pre>

<h3>步驟三：發送 API 請求</h3>
<pre><span class="code-label">PYTHON</span><code class="language-python">import requests, json

params["sign"] = sign

url = "https://uat.fintechpaymenthub.com/openapi/v1/pay"
response = requests.post(url, json=params, headers={"Content-Type": "application/json"})
result = response.json()

if result.get("message") == "SUCCESS":
    print("下單成功！", result)
else:
    print("下單失敗：", result)</code></pre>

<h3>步驟四：處理支付通知</h3>
<p>支付完成後，平台會向 <code>notify_url</code> 發送 POST 請求通知支付結果。商戶需：</p>
<ul>
  <li>驗證簽名是否正確</li>
  <li>校驗訂單號和金額與本地系統一致</li>
  <li>處理業務邏輯後返回純字符串 <code>success</code></li>
</ul>

<h2 id="code-example">完整代碼示例</h2>
<p>以下是一個完整的 Python 示例，涵蓋下單、查詢和通知處理：</p>
<pre><span class="code-label">PYTHON</span><code class="language-python">import hashlib, json, requests
from flask import Flask, request

app = Flask(__name__)
MERCHANT_KEY = "YOUR_MERCHANT_KEY"
BASE_URL = "https://uat.fintechpaymenthub.com/openapi/v1"

def sign(params, key=MERCHANT_KEY):
    sorted_params = sorted(params.items())
    sign_str = "&".join(f"{k}={v}" for k, v in sorted_params if v)
    sign_str += f"&key={key}"
    return hashlib.md5(sign_str.encode("utf-8")).hexdigest().upper()

# 下單
def create_order():
    params = {
        "service": "pay.weixin.web.intl",
        "client_id": "1001", "store_id": "1001001",
        "order_no": "ORDER20260708001", "amount": "1.00",
        "body": "測試商品", "notify_url": "https://your-domain.com/notify",
        "callback_url": "https://your-domain.com/result",
        "client_ip": "127.0.0.1", "nonce_str": "abc123def456",
    }
    params["sign"] = sign(params)
    resp = requests.post(f"{BASE_URL}/pay", json=params)
    return resp.json()

# 查詢訂單
def query_order(order_no):
    params = {
        "service": "pay.weixin.web.intl",
        "client_id": "1001", "store_id": "1001001",
        "order_no": order_no, "nonce_str": "query123456",
    }
    params["sign"] = sign(params)
    resp = requests.post(f"{BASE_URL}/query", json=params)
    return resp.json()

# 接收支付通知
@app.route("/notify", methods=["POST"])
def notify():
    data = request.get_json()
    # 驗證簽名
    recv_sign = data.pop("sign", "")
    if sign(data) != recv_sign:
        return "FAIL"
    # 校驗訂單
    if data.get("message") == "SUCCESS":
        # 更新本地訂單狀態
        print(f"訂單 {data.get('order_no')} 支付成功")
    return "success"</code></pre>
<p>更多語言示例請參考 <a href="signature.html">數字簽名</a> 頁面。</p>
'''
    _write_common_page('quick-start', '快速入門', content, toc)

    # 2. API Endpoints
    toc = [
        {'level': 2, 'text': '環境地址', 'slug': 'environments'},
        {'level': 2, 'text': '端點摘要', 'slug': 'endpoint-summary'},
        {'level': 2, 'text': '各支付方式接口支持', 'slug': 'interface-matrix'},
    ]
    matrix_rows = ''
    visible_doc_ids = {d['id'] for d in DOCS}
    for doc_id, name, supports in _INTERFACE_MATRIX:
        if doc_id not in visible_doc_ids:
            continue
        cat = next((d['category'] for d in DOCS if d['id'] == doc_id), '')
        link = f'<a href="../payments/{cat}/{doc_id}.html">{name}</a>'
        matrix_rows += f'<tr><td>{link}</td>'
        for s in supports:
            cls = 'param-required yes' if s == '✅' else 'param-required no'
            matrix_rows += f'<td><span class="{cls}">{s}</span></td>'
        matrix_rows += '</tr>\n'

    content = f'''<h1 id="overview">API 調用地址</h1>
<div class="page-meta">
  <span class="page-meta-item">更新日期：2026-03-13</span>
</div>
<div class="callout info">
  <div class="callout-title">ℹ️ 基礎 URL</div>
  <p>所有 API 請求均需以以下基礎 URL 為前綴，再加上各接口特定的路徑。</p>
</div>

<h2 id="environments">環境地址</h2>
<div class="table-wrapper">
<table>
<thead><tr><th>環境</th><th>基礎 URL</th></tr></thead>
<tbody>
<tr><td>UAT（測試環境）</td><td><code>https://uat.fintechpaymenthub.com/openapi/v1/</code></td></tr>
<tr><td>Production（生產環境）</td><td><code>https://api.fintechpaymenthub.com/openapi/v1/</code></td></tr>
</tbody>
</table>
</div>

<h2 id="endpoint-summary">端點摘要</h2>
<p>絕大多數渠道的<strong>下單、查詢、退款</strong>功能都使用相同的三個端點，僅透過 <code>service</code> 參數區分具體支付渠道。</p>
<div class="table-wrapper">
<table>
<thead><tr><th>功能</th><th>端點路徑</th><th>請求方法</th><th>說明</th></tr></thead>
<tbody>
<tr><td>下單</td><td><code>/pay</code></td><td><span class="param-required yes">POST</span></td><td>創建支付訂單</td></tr>
<tr><td>調起支付</td><td><code>/pay-again</code></td><td><span class="param-required yes">POST</span></td><td>重新獲取支付鏈接（部分渠道支持）</td></tr>
<tr><td>訂單查詢</td><td><code>/query</code></td><td><span class="param-required yes">POST</span></td><td>查詢訂單狀態</td></tr>
<tr><td>退款</td><td><code>/refund</code></td><td><span class="param-required yes">POST</span></td><td>發起退款</td></tr>
<tr><td>退款查詢</td><td><code>/refundQuery</code></td><td><span class="param-required yes">POST</span></td><td>查詢退款狀態</td></tr>
</tbody>
</table>
</div>

<h2 id="interface-matrix">各支付方式接口支持</h2>
<p>不同支付方式支持的接口有所不同，請參考下表：</p>
<div class="table-wrapper">
<table>
<thead><tr><th>支付方式</th><th>下單</th><th>調起支付</th><th>查詢</th><th>退款</th><th>退款查詢</th><th>支付通知</th></tr></thead>
<tbody>
{matrix_rows}
</tbody>
</table>
</div>
'''
    _write_common_page('endpoints', 'API 調用地址', content, toc)

    # 3. Data Format
    toc = [
        {'level': 2, 'text': '請求格式', 'slug': 'request-format'},
        {'level': 2, 'text': '響應格式', 'slug': 'response-format'},
        {'level': 2, 'text': '公共請求參數', 'slug': 'common-params'},
    ]
    content = '''<h1 id="overview">數據格式</h1>
<div class="page-meta">
  <span class="page-meta-item">文檔版本：V1.0</span>
</div>
<p>本頁面說明 Fintech Payment OpenAPI 的數據傳輸格式，包括請求格式、響應格式及公共參數。</p>

<h2 id="request-format">請求格式</h2>
<p>API 支持兩種請求提交方式：</p>

<h3>標準模式（JSON POST）</h3>
<p>絕大多數支付渠道採用 HTTPS POST + JSON 格式提交數據：</p>
<ul>
  <li>請求方法：<code>POST</code></li>
  <li>Content-Type：<code>application/json</code></li>
  <li>字符編碼：<code>UTF-8</code></li>
  <li>數據必須簽名</li>
</ul>
<pre><span class="code-label">JSON</span><code class="language-json">{
  "service": "pay.weixin.web.intl",
  "order_no": "00f78e479fe74a11a73be0feb9999749",
  "body": "測試支付",
  "amount": 1.00,
  "callback_url": "https://your-domain.com/result",
  "notify_url": "https://your-domain.com/notify",
  "client_ip": "127.0.0.1",
  "client_id": "1002",
  "store_id": "1001002",
  "nonce_str": "ptf3tyjv",
  "sign": "64950ED7B13489114C671EFD24C45AC3",
  "lang": "en_US"
}</code></pre>

<h3>託管模式（FORM POST）</h3>
<p>Global Payments Online 託管模式使用 HTML 表單提交，適合直接在瀏覽器中跳轉到支付頁面：</p>
<ul>
  <li>請求方法：<code>POST</code></li>
  <li>Content-Type：<code>application/x-www-form-urlencoded</code></li>
  <li>字符編碼：<code>UTF-8</code></li>
  <li>數據必須簽名</li>
</ul>
<pre><span class="code-label">HTML</span><code class="language-html">&lt;form action="https://uat.fintechpaymenthub.com/openapi/v1/pay" method="post"&gt;
  &lt;input type="hidden" name="service" value="pay.globalpayment.card.notpresent"&gt;
  &lt;input type="hidden" name="client_id" value="1001"&gt;
  &lt;input type="hidden" name="store_id" value="1001001"&gt;
  &lt;input type="hidden" name="order_no" value="ORDER001"&gt;
  &lt;input type="hidden" name="amount" value="100"&gt;
  &lt;input type="hidden" name="sign" value="XXXXX"&gt;
&lt;/form&gt;</code></pre>

<h2 id="response-format">響應格式</h2>
<p>所有接口返回 JSON 格式數據。根據 <code>message</code> 字段判斷請求是否成功：</p>

<h3>成功返回</h3>
<pre><span class="code-label">JSON</span><code class="language-json">{
  "message": "SUCCESS",
  "order_no": "00f78e479fe74a11a73be0feb9999749",
  "transaction_id": "FTP202012181649370001",
  "submit_time": "20201218164937",
  "nonce_str": "o66w9pft",
  "sign": "3F6C05280C90E5BB1330524FD61C6121"
}</code></pre>
<p><code>message</code> 返回 <code>SUCCESS</code> 表示接口調用成功。返回結果中包含 <code>sign</code> 字段，商戶應驗證簽名以確保數據完整性。</p>

<h3>錯誤返回</h3>
<pre><span class="code-label">JSON</span><code class="language-json">{
  "message": "Missing Merchant ID",
  "errors": [
    {
      "field": "mch_id",
      "message": "Missing Merchant ID",
      "code": "MISSING_FIELD"
    }
  ]
}</code></pre>
<p><code>message</code> 返回非 <code>SUCCESS</code> 時表示接口調用出錯。<code>errors</code> 數組包含具體的錯誤字段和錯誤碼。</p>

<h2 id="common-params">公共請求參數</h2>
<p>以下參數在所有接口中通用：</p>
<div class="table-wrapper">
<table>
<thead><tr><th>參數名</th><th>類型</th><th>必填</th><th>說明</th><th>是否參與簽名</th></tr></thead>
<tbody>
<tr><td><code>service</code></td><td>String(40)</td><td><span class="param-required yes">必填</span></td><td>支付服務標識，區分不同支付渠道</td><td>✅</td></tr>
<tr><td><code>client_id</code></td><td>String(32)</td><td><span class="param-required yes">必填</span></td><td>商戶客戶端 ID</td><td>✅</td></tr>
<tr><td><code>store_id</code></td><td>String(50)</td><td><span class="param-required yes">必填</span></td><td>門店 ID</td><td>✅</td></tr>
<tr><td><code>order_no</code></td><td>String(32)</td><td><span class="param-required yes">必填</span></td><td>商戶系統內部的訂單號，需確保唯一</td><td>✅</td></tr>
<tr><td><code>nonce_str</code></td><td>String(32)</td><td><span class="param-required yes">必填</span></td><td>隨機字符串，防重放攻擊</td><td>✅</td></tr>
<tr><td><code>sign</code></td><td>String(64)</td><td><span class="param-required yes">必填</span></td><td>數字簽名，詳見<a href="signature.html">數字簽名</a></td><td>❌</td></tr>
<tr><td><code>lang</code></td><td>String(10)</td><td><span class="param-required no">選填</span></td><td>語言設置，如 <code>en_US</code>、<code>zh_CN</code></td><td>❌</td></tr>
</tbody>
</table>
</div>
'''
    _write_common_page('data-format', '數據格式', content, toc)

    # 4. Digital Signature
    toc = [
        {'level': 2, 'text': '簽名原始串', 'slug': 'sign-string'},
        {'level': 2, 'text': '簽名演算法', 'slug': 'sign-algorithm'},
        {'level': 2, 'text': '完整示例', 'slug': 'full-example'},
        {'level': 2, 'text': '多語言代碼示例', 'slug': 'code-examples'},
    ]
    content = '''<h1 id="overview">數字簽名</h1>
<div class="page-meta">
  <span class="page-meta-item">文檔版本：V1.0</span>
</div>
<p>為保證資料傳輸過程中的真實性和完整性，所有 API 請求和響應都需要進行數字簽名。簽名分為兩步：先按規則拼接簽名原始串，再用 MD5 演算法計算簽名結果。</p>

<h2 id="sign-string">簽名原始串</h2>
<p>無論是請求還是應答，簽名原始串按以下規則組裝：</p>
<ol>
  <li>除 <code>sign</code> 和 <code>lang</code> 字段外，所有參數按照字段名的 ASCII 碼從小到大排序，使用 QueryString 格式（即 <code>key1=value1&amp;key2=value2&amp;...</code>）拼接</li>
  <li>空值不傳遞，不參與簽名組串</li>
  <li>字段名和字段值都採用原始值，不進行 URL Encode</li>
  <li>平台返回的應答或通知可能會由於升級增加參數，請驗證應答簽名時注意允許這種情況</li>
</ol>

<h3>示例</h3>
<p>假設有以下請求參數：</p>
<pre><span class="code-label">JSON</span><code class="language-json">{
  "client_id": "1001",
  "store_id": "1001001",
  "service": "pay.weixin.web.intl",
  "order_no": "Me08772e914454656bf26ec8996e35bb",
  "amount": "1",
  "body": "a chair",
  "notify_url": "2031",
  "client_ip": "127.0.0.1",
  "nonce_str": "1409196838",
  "sign": "22C9835251DCF58BB70C78FA8D6448BD"
}</code></pre>
<p>正確的簽名排序結果為（排除 <code>sign</code> 和 <code>lang</code>）：</p>
<pre><span class="code-label">TEXT</span><code class="language-text">amount=1&body=a chair&client_ip=127.0.0.1&client_id=1001&nonce_str=1409196838&notify_url=2031&order_no=Me08772e914454656bf26ec8996e35bb&service=pay.weixin.web.intl&store_id=1001001</code></pre>

<h2 id="sign-algorithm">簽名演算法</h2>
<p>使用 MD5 摘要演算法，在簽名原始串後加上商戶通信金鑰進行計算，結果轉為大寫：</p>
<div class="callout info">
  <div class="callout-title">📐 簽名公式</div>
  <p><code>sign = MD5(簽名原始串 + "&key=" + 商戶秘鑰).toUpperCase()</code></p>
</div>
<p>注意：簽名時將字符串轉化為位元組流時指定的編碼字符集應為 <code>UTF-8</code>。</p>

<h2 id="full-example">完整示例</h2>
<p>假設商戶秘鑰為 <code>MjNiYWE4NGY2YTMxNDk2MTljODBkMjZkMjk0NjAxY2M=</code>：</p>
<pre><span class="code-label">TEXT</span><code class="language-text">簽名原始串：
amount=1&body=a chair&client_ip=127.0.0.1&client_id=1001&nonce_str=1409196838&notify_url=2031&order_no=Me08772e914454656bf26ec8996e35bb&service=pay.weixin.web.intl&store_id=1001001

拼接金鑰後：
amount=1&body=a chair&client_ip=127.0.0.1&client_id=1001&nonce_str=1409196838&notify_url=2031&order_no=Me08772e914454656bf26ec8996e35bb&service=pay.weixin.web.intl&store_id=1001001&key=MjNiYWE4NGY2YTMxNDk2MTljODBkMjZkMjk0NjAxY2M=

MD5 運算並轉大寫：
52A4A2D2ACA971B309A0B9B1F21C8698</code></pre>

<h2 id="code-examples">多語言代碼示例</h2>

<h3>Python</h3>
<pre><span class="code-label">PYTHON</span><code class="language-python">import hashlib

def generate_sign(params, merchant_key):
    # 排除 sign 和 lang，按 key ASCII 碼排序
    filtered = {k: v for k, v in params.items()
                if k not in ("sign", "lang") and v}
    sorted_params = sorted(filtered.items())
    # 拼接成 key=value&key=value
    sign_str = "&".join(f"{k}={v}" for k, v in sorted_params)
    # 拼接商戶秘鑰
    sign_str += f"&key={merchant_key}"
    # MD5 轉大寫
    return hashlib.md5(sign_str.encode("utf-8")).hexdigest().upper()</code></pre>

<h3>Java</h3>
<pre><span class="code-label">JAVA</span><code class="language-java">import java.security.MessageDigest;
import java.util.*;

public String generateSign(Map&lt;String, String&gt; params, String merchantKey) throws Exception {
    // 排除 sign 和 lang
    params.remove("sign");
    params.remove("lang");
    // 按 key 排序
    List&lt;String&gt; keys = new ArrayList&lt;&gt;(params.keySet());
    Collections.sort(keys);
    // 拼接
    StringBuilder sb = new StringBuilder();
    for (String key : keys) {
        String val = params.get(key);
        if (val != null && !val.isEmpty()) {
            if (sb.length() > 0) sb.append("&amp;");
            sb.append(key).append("=").append(val);
        }
    }
    sb.append("&amp;key=").append(merchantKey);
    // MD5
    MessageDigest md = MessageDigest.getInstance("MD5");
    byte[] digest = md.digest(sb.toString().getBytes("UTF-8"));
    StringBuilder hex = new StringBuilder();
    for (byte b : digest) hex.append(String.format("%02X", b));
    return hex.toString();
}</code></pre>

<h3>PHP</h3>
<pre><span class="code-label">PHP</span><code class="language-php">&lt;?php
function generateSign($params, $merchantKey) {
    // 排除 sign 和 lang
    unset($params['sign'], $params['lang']);
    // 按 key 排序
    ksort($params);
    // 拼接
    $parts = [];
    foreach ($params as $k =&gt; $v) {
        if ($v !== '' &amp;&amp; $v !== null) {
            $parts[] = "{$k}={$v}";
        }
    }
    $signStr = implode('&amp;', $parts) . "&amp;key={$merchantKey}";
    // MD5 轉大寫
    return strtoupper(md5($signStr));
}</code></pre>

<h3>Node.js</h3>
<pre><span class="code-label">JAVASCRIPT</span><code class="language-javascript">const crypto = require("crypto");

function generateSign(params, merchantKey) {
    // 排除 sign 和 lang
    const { sign, lang, ...rest } = params;
    // 按 key 排序
    const sortedKeys = Object.keys(rest).sort();
    // 拼接
    const signStr = sortedKeys
        .filter(k => rest[k])
        .map(k => `${k}=${rest[k]}`)
        .join("&") + `&key=${merchantKey}`;
    // MD5 轉大寫
    return crypto.createHash("md5").update(signStr, "utf8").digest("hex").toUpperCase();
}</code></pre>
'''
    _write_common_page('signature', '數字簽名', content, toc)

    # 5. Request Headers
    toc = [
        {'level': 2, 'text': '標準請求頭', 'slug': 'standard-headers'},
        {'level': 2, 'text': '託管模式請求頭', 'slug': 'hosted-headers'},
    ]
    content = '''<h1 id="overview">請求頭說明</h1>
<div class="page-meta">
  <span class="page-meta-item">文檔版本：V1.0</span>
</div>
<p>本頁面說明調用 Fintech Payment OpenAPI 時所需的 HTTP 請求頭。</p>

<h2 id="standard-headers">標準請求頭</h2>
<p>使用 JSON POST 方式調用接口時，請設置以下請求頭：</p>
<div class="table-wrapper">
<table>
<thead><tr><th>請求頭</th><th>值</th><th>必填</th><th>說明</th></tr></thead>
<tbody>
<tr><td><code>Content-Type</code></td><td><code>application/json</code></td><td><span class="param-required yes">必填</span></td><td>請求體為 JSON 格式</td></tr>
<tr><td><code>Accept</code></td><td><code>application/json</code></td><td><span class="param-required no">建議</span></td><td>響應體為 JSON 格式</td></tr>
<tr><td><code>Accept-Charset</code></td><td><code>utf-8</code></td><td><span class="param-required no">建議</span></td><td>字符編碼為 UTF-8</td></tr>
</tbody>
</table>
</div>

<h2 id="hosted-headers">託管模式請求頭</h2>
<p>Global Payments Online 託管模式使用 HTML 表單提交，請求頭有所不同：</p>
<div class="table-wrapper">
<table>
<thead><tr><th>請求頭</th><th>值</th><th>說明</th></tr></thead>
<tbody>
<tr><td><code>Content-Type</code></td><td><code>application/x-www-form-urlencoded</code></td><td>表單提交格式</td></tr>
</tbody>
</table>
</div>

<div class="callout info">
  <div class="callout-title">ℹ️ 字符編碼</div>
  <p>所有請求和響應均使用 <strong>UTF-8</strong> 編碼。簽名計算時也必須使用 UTF-8 編碼將字符串轉換為位元組流。</p>
</div>
'''
    _write_common_page('request-headers', '請求頭說明', content, toc)

    # 6. Error Codes
    toc = [
        {'level': 2, 'text': '錯誤響應結構', 'slug': 'error-structure'},
        {'level': 2, 'text': '通用錯誤碼', 'slug': 'common-errors'},
        {'level': 2, 'text': '卡支付專屬錯誤碼', 'slug': 'card-errors'},
    ]
    common_rows = ''.join(f'<tr><td><code>{c}</code></td><td>{d}</td></tr>\n' for c, d in _COMMON_ERRORS)
    card_rows = ''.join(f'<tr><td><code>{c}</code></td><td>{d}</td></tr>\n' for c, d in _CARD_ERRORS)
    content = f'''<h1 id="overview">錯誤碼說明</h1>
<div class="page-meta">
  <span class="page-meta-item">文檔版本：V1.0</span>
</div>
<p>本頁面列出 Fintech Payment OpenAPI 可能返回的所有錯誤碼，分為通用錯誤碼和卡支付專屬錯誤碼。</p>

<h2 id="error-structure">錯誤響應結構</h2>
<p>當接口調用出錯時，返回以下 JSON 結構：</p>
<pre><span class="code-label">JSON</span><code class="language-json">{{
  "message": "錯誤描述信息",
  "errors": [
    {{
      "field": "出錯字段名",
      "message": "字段錯誤描述",
      "code": "錯誤碼"
    }}
  ]
}}</code></pre>
<p><code>message</code> 返回非 <code>SUCCESS</code> 時表示接口調用出錯。</p>

<h2 id="common-errors">通用錯誤碼</h2>
<p>以下錯誤碼在所有支付渠道中通用（共 {len(_COMMON_ERRORS)} 個）：</p>
<div class="table-wrapper">
<table>
<thead><tr><th>錯誤碼</th><th>說明</th></tr></thead>
<tbody>
{common_rows}
</tbody>
</table>
</div>

<h2 id="card-errors">卡支付專屬錯誤碼</h2>
<p>以下錯誤碼僅在 Card Pay（Visa/Mastercard/銀聯等卡支付）場景中出現（共 {len(_CARD_ERRORS)} 個）：</p>
<div class="table-wrapper">
<table>
<thead><tr><th>錯誤碼</th><th>說明</th></tr></thead>
<tbody>
{card_rows}
</tbody>
</table>
</div>
'''
    _write_common_page('error-codes', '錯誤碼說明', content, toc)

    # 7. Payment Notification
    toc = [
        {'level': 2, 'text': '通知機制', 'slug': 'mechanism'},
        {'level': 2, 'text': '重試策略', 'slug': 'retry-strategy'},
        {'level': 2, 'text': '處理建議', 'slug': 'best-practices'},
    ]
    content = '''<h1 id="overview">支付通知說明</h1>
<div class="page-meta">
  <span class="page-meta-item">文檔版本：V1.0</span>
</div>
<p>支付完成後，平台會異步通知商戶支付結果。本頁面說明通知機制、重試策略及處理建議。</p>

<h2 id="mechanism">通知機制</h2>
<p>支付完成後，系統會將支付結果通知到下單時填寫的 <code>notify_url</code>：</p>
<ul>
  <li>通知方式：<code>POST</code> 請求，數據體為 JSON 格式</li>
  <li>通知 URL：下單接口中提交的 <code>notify_url</code> 參數</li>
  <li>通知內容包含訂單號、交易號、支付金額、支付狀態等信息，並附帶 <code>sign</code> 簽名</li>
</ul>

<h2 id="retry-strategy">重試策略</h2>
<p>如果平台收到商戶的應答不是純字符串 <code>success</code>，或超過 5 秒未返回，平台認為通知失敗，將按以下策略重新發起通知：</p>
<div class="table-wrapper">
<table>
<thead><tr><th>重試次數</th><th>間隔時間</th></tr></thead>
<tbody>
<tr><td>第 1 次</td><td>0 秒（立即）</td></tr>
<tr><td>第 2 次</td><td>15 秒</td></tr>
<tr><td>第 3 次</td><td>15 秒</td></tr>
<tr><td>第 4 次</td><td>30 秒</td></tr>
<tr><td>第 5 次</td><td>180 秒（3 分鐘）</td></tr>
<tr><td>第 6 次</td><td>1800 秒（30 分鐘）</td></tr>
<tr><td>第 7-9 次</td><td>1800 秒（30 分鐘）</td></tr>
<tr><td>第 10 次</td><td>3600 秒（1 小時）</td></tr>
</tbody>
</table>
</div>
<div class="callout warning">
  <div class="callout-title">⚠️ 注意</div>
  <p>平台會盡可能提高通知的成功率，但不保證通知最終能成功。建議商戶同時使用<strong>主動查詢</strong>（調用 <code>/query</code> 接口）作為兜底方案。</p>
</div>

<h2 id="best-practices">處理建議</h2>
<p>由於存在重新發送後台通知的情況，同樣的通知可能會多次發送給商戶系統。商戶系統必須能夠正確處理重複通知：</p>
<ol>
  <li><strong>驗證簽名</strong>：收到通知後，先驗證 <code>sign</code> 字段是否正確，防止偽造</li>
  <li><strong>校驗訂單</strong>：校驗通知中的 <code>order_no</code> 和 <code>amount</code> 與自身業務系統的訂單和金額是否一致</li>
  <li><strong>冪等處理</strong>：檢查對應訂單是否已處理過，如已處理直接返回 <code>success</code></li>
  <li><strong>併發控制</strong>：採用數據鎖進行併發控制，避免函數重入造成數據混亂</li>
  <li><strong>及時應答</strong>：處理完成後立即返回純字符串 <code>success</code>（5 秒內），避免觸發重試</li>
</ol>

<h3>通知處理代碼示例</h3>
<pre><span class="code-label">PYTHON</span><code class="language-python">@app.route("/notify", methods=["POST"])
def handle_notify():
    data = request.get_json()

    # 1. 驗證簽名
    recv_sign = data.pop("sign", "")
    if generate_sign(data) != recv_sign:
        return "FAIL"

    # 2. 校驗訂單
    order_no = data.get("order_no")
    amount = data.get("amount")
    order = get_order_from_db(order_no)
    if not order or float(order.amount) != float(amount):
        return "FAIL"

    # 3. 冪等處理
    if order.status == "PAID":
        return "success"  # 已處理，直接返回成功

    # 4. 更新訂單狀態（加鎖防併發）
    with db_lock:
        update_order_status(order_no, "PAID")

    # 5. 返回成功
    return "success"</code></pre>
'''
    _write_common_page('notification', '支付通知說明', content, toc)


def build_homepage():
    """Build the homepage with payment channel overview."""
    print("\n🏠 建構首頁...")

    # Group docs by category
    grouped = {}
    for d in DOCS:
        cat = d['category']
        if cat not in grouped:
            grouped[cat] = []
        grouped[cat].append(d)

    cards_html = ''
    for cat_key in CATEGORY_ORDER:
        if cat_key not in grouped:
            continue
        cat_info = CATEGORIES[cat_key]
        docs_list = grouped[cat_key]

        cards_html += f'''
<div class="category-card {cat_info['color']}">
  <div class="icon">{cat_info['icon']}</div>
  <h3>{cat_info['label']}</h3>
  <p>{cat_info['desc']}</p>
  <ul class="links">
'''
        for d in docs_list:
            status_badge = ''
            if d['status'] == 'upcoming':
                status_badge = ' <span class="status-tag upcoming">暫未開放</span>'
            cards_html += f'    <li><a href="payments/{cat_key}/{d["id"]}.html">{html.escape(d["title"])}{status_badge}<span class="arrow">→</span></a></li>\n'
        cards_html += '  </ul>\n</div>\n'

    content_html = f'''
<div class="home-hero">
  <h1>Fintech Payment OpenAPI</h1>
  <p>為您提供安全、便捷、多渠道的支付接口服務，支持卡支付、行動支付、掃碼支付等多種支付方式。</p>
  <div class="cta-buttons">
    <a href="common/quick-start.html" class="btn btn-primary">🚀 開始接入</a>
    <a href="common/data-format.html" class="btn btn-secondary">📚 查看文檔</a>
  </div>
</div>

<section class="payment-categories">
  <h2>支付渠道一覽</h2>
  <div class="category-grid">
    {cards_html}
  </div>
</section>

<section style="padding: 40px 0;">
  <div class="callout info">
    <div class="callout-title">ℹ️ 接入前提醒</div>
    <p>所有接口調用前，請先閱讀 <a href="common/quick-start.html">快速入門</a> 和 <a href="common/data-format.html">通用規則</a>，確保接口調用正確。</p>
  </div>
</section>
'''

    nav_html = generate_sidebar('', '', depth=0)
    first_doc = DOCS[0]

    full_html = f'''<!DOCTYPE html>
<html lang="zh-Hant">
<head>
  <meta charset="UTF-8">
  <meta name="viewport" content="width=device-width, initial-scale=1.0">
  <title>Fintech Payment OpenAPI 文檔</title>
  <link rel="stylesheet" href="assets/css/style.css">
</head>
<body>
  <header class="site-header">
    <div class="header-inner">
      <a href="index.html" class="logo">
        <span class="logo-icon">FP</span>
        <span class="logo-text">Fintech Payment <span class="logo-sub">OpenAPI</span></span>
      </a>
      <nav class="header-nav">
        <a href="index.html" class="active">首頁</a>
        <a href="common/quick-start.html">接入指引</a>
        <a href="common/data-format.html">通用規則</a>
        <a href="payments/{first_doc["category"]}/{first_doc["id"]}.html">產品文檔</a>
      </nav>
      <button class="menu-toggle" aria-label="選單">☰</button>
    </div>
  </header>

  <div class="docs-layout">
    {nav_html}

    <main class="main-content" style="max-width: 1100px;">
      {content_html}

      <footer class="site-footer">
        <p>&copy; 2026 Fintech Payment. 保留所有權利。 | <a href="https://fintechpayment.com" target="_blank">官方網站</a></p>
      </footer>
    </main>
  </div>

  <script src="assets/js/main.js"></script>
</body>
</html>'''

    output_path = os.path.join(OUTPUT_DIR, 'index.html')
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(full_html)


def cleanup_old_output():
    """Remove old category directories that are no longer used."""
    old_cats = ['card', 'online', 'offline', 'special']
    for old_cat in old_cats:
        old_path = os.path.join(OUTPUT_DIR, 'payments', old_cat)
        if os.path.exists(old_path):
            shutil.rmtree(old_path)
            print(f"  🗑️ 清理舊目錄: payments/{old_cat}/")


def load_visibility_config():
    """Load docs-config.json and return the hidden doc IDs list."""
    config_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'docs-config.json')
    hidden = []
    if os.path.exists(config_path):
        try:
            with open(config_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
                hidden = data.get('hidden', [])
        except (json.JSONDecodeError, Exception) as e:
            print(f"  ⚠️ 配置文件讀取錯誤: {e}，忽略配置")
    return hidden


def apply_visibility_filter(hidden_ids):
    """Filter DOCS and CATEGORY_ORDER based on hidden IDs.
    Returns (visible_docs, visible_categories).
    """
    global DOCS, CATEGORY_ORDER
    if not hidden_ids:
        return DOCS, CATEGORY_ORDER

    # Filter out hidden docs
    visible_docs = [d for d in DOCS if d['id'] not in hidden_ids]
    if len(visible_docs) < len(DOCS):
        skipped = [d['title'] for d in DOCS if d['id'] in hidden_ids]
        print(f"\n🔇 隱藏文檔: {', '.join(skipped)}")

    # Filter categories — only keep categories that have at least one visible doc
    visible_cats = []
    for cat in CATEGORY_ORDER:
        if any(d['category'] == cat for d in visible_docs):
            visible_cats.append(cat)
        else:
            cat_label = CATEGORIES[cat]['label'] if cat in CATEGORIES else cat
            print(f"  🗑️ 隱藏空分類: {cat_label}")

    return visible_docs, visible_cats


def main():
    global DOCS, CATEGORY_ORDER

    print("🚀 Fintech Payment OpenAPI 文檔網站生成器")
    print("=" * 50)

    # Load visibility config (unless --show-all is specified)
    show_all = '--show-all' in sys.argv
    if show_all:
        print("\n📋 --show-all: 顯示所有文檔（忽略隱藏配置）")
        docs_to_process = list(DOCS)
        cat_order = list(CATEGORY_ORDER)
    else:
        hidden_ids = load_visibility_config()
        docs_to_process, cat_order = apply_visibility_filter(hidden_ids)

    # Override globals for all downstream functions
    DOCS = docs_to_process
    CATEGORY_ORDER = cat_order

    # Clean up old output directories
    print("\n🧹 清理舊輸出...")
    cleanup_old_output()

    # Build all payment docs
    print("\n📄 處理支付產品文檔...")
    success_count = 0
    for doc_info in DOCS:
        if build_document(doc_info):
            success_count += 1

    print(f"\n✅ 成功處理 {success_count}/{len(DOCS)} 份文檔")

    # Build common pages
    build_common_pages()

    # Build homepage
    build_homepage()

    print("\n🎉 完成！文檔網站已生成。")
    print(f"📁 輸出目錄: {OUTPUT_DIR}")


if __name__ == '__main__':
    main()
